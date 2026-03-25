"""
WorkanaBot — Servidor Python (Flask)
"""

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from functools import wraps
import asyncio
import copy
import json
import os
import re
import subprocess
import sys
import threading
import time
import uuid
import logging

from dotenv import load_dotenv
load_dotenv()

# ── CONFIGURAÇÃO ──────────────────────────────────────────────────────────────

logging.basicConfig(level=logging.INFO, format="[WorkanaBot Server] %(message)s")
log = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app, resources={r"/api/*": {"origins": "*"}})

SERVER_PORT = 5002
DATA_DIR    = os.path.join(os.path.dirname(__file__), "data")
os.makedirs(DATA_DIR, exist_ok=True)

SCRAPER_PATH = os.path.join(os.path.dirname(__file__), "scraper.py")

# ── SESSION KEEPER ────────────────────────────────────────────────────────────

SAME_SITE_MAP = {
    "strict": "Strict", "lax": "Lax", "none": "None",
    "no_restriction": "None", "unspecified": "Lax", "": "Lax",
}

def _normalizar_cookies(cookies: list) -> list:
    result = []
    for c in cookies:
        c = dict(c)
        c["sameSite"] = SAME_SITE_MAP.get(str(c.get("sameSite") or "").lower(), "Lax")
        for k in ("expirationDate", "hostOnly", "session", "storeId"):
            c.pop(k, None)
        result.append(c)
    return result


class SessionKeeper:
    """Mantém um browser headless aberto e renova a sessão Workana a cada 30 min."""

    RENOVAR_INTERVALO = 900  # 15 minutos

    def __init__(self):
        self._thread: threading.Thread | None = None
        self._loop: asyncio.AbstractEventLoop | None = None
        self._cookies: list = []
        self._running = False

    def start(self, cookies: list):
        self._cookies = _normalizar_cookies(cookies)
        if self._thread and self._thread.is_alive():
            return
        self._running = True
        self._thread = threading.Thread(target=self._run, daemon=True, name="session-keeper")
        self._thread.start()
        log.info("SessionKeeper: iniciado.")

    def _run(self):
        self._loop = asyncio.new_event_loop()
        asyncio.set_event_loop(self._loop)
        self._loop.run_until_complete(self._keeper_loop())

    async def _keeper_loop(self):
        try:
            from playwright.async_api import async_playwright
        except ImportError:
            log.warning("SessionKeeper: playwright não disponível.")
            return

        async with async_playwright() as pw:
            browser = await pw.chromium.launch(headless=True)
            context = await browser.new_context(
                user_agent=(
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) "
                    "Chrome/124.0.0.0 Safari/537.36"
                ),
                locale="pt-BR",
            )
            await context.add_cookies(self._cookies)
            log.info("SessionKeeper: browser headless ativo.")

            while self._running:
                try:
                    page = await context.new_page()
                    await page.goto(
                        "https://www.workana.com/projects",
                        wait_until="domcontentloaded",
                        timeout=20000,
                    )
                    fresh = await context.cookies()
                    await page.close()

                    if estado.get("sessao"):
                        estado["sessao"]["cookies"] = fresh
                        estado["sessao"]["salvo_em"] = int(time.time())
                        _salvar_json("sessao.json", estado["sessao"])
                        log.info(f"SessionKeeper: sessão renovada ({len(fresh)} cookies).")
                except Exception as e:
                    log.warning(f"SessionKeeper: erro ao renovar sessão: {e}")

                await asyncio.sleep(self.RENOVAR_INTERVALO)
                # Dispara scraping para buscar novas vagas
                if estado.get("sessao"):
                    iniciar_scraping_async(estado["sessao"].get("cookies", []))
                    log.info("SessionKeeper: scraping automático disparado.")

            await browser.close()


session_keeper = SessionKeeper()

# ── ESTADO EM MEMÓRIA ─────────────────────────────────────────────────────────

estado = {
    "sessao":              None,
    "projetos":            [],
    "projetos_salvos_em":  0,
    "scraping_em_curso":   False,
    "ultimo_scraping":     0,
    "scraping_erro":       None,
    "projetos_salvos":     [],
}

# ── UTILITÁRIOS DE DISCO ──────────────────────────────────────────────────────

def _salvar_json(nome, dados):
    try:
        with open(os.path.join(DATA_DIR, nome), "w", encoding="utf-8") as f:
            json.dump(dados, f, ensure_ascii=False, indent=2)
    except Exception as e:
        log.error(f"Erro ao salvar {nome}: {e}")

def _carregar_json(nome):
    caminho = os.path.join(DATA_DIR, nome)
    if not os.path.exists(caminho):
        return None
    try:
        with open(caminho, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        log.error(f"Erro ao carregar {nome}: {e}")
        return None

# Bootstrapping
log.info("Carregando cache do disco...")
estado["sessao"]   = _carregar_json("sessao.json")
cache_proj         = _carregar_json("projetos.json")
if cache_proj:
    estado["projetos"]           = cache_proj.get("projetos", [])
    estado["projetos_salvos_em"] = cache_proj.get("salvo_em", 0)

proj_salvos = _carregar_json("projetos_salvos.json")
if proj_salvos:
    estado["projetos_salvos"] = proj_salvos if isinstance(proj_salvos, list) else []

PROPOSTAS_DIR = os.path.join(DATA_DIR, "propostas")
os.makedirs(PROPOSTAS_DIR, exist_ok=True)

# ── AUTH ──────────────────────────────────────────────────────────────────────

APP_TOKEN = os.environ.get("WORKANABOT_TOKEN", "workanabot-dev-token")

def requer_auth(f):
    @wraps(f)
    def wrapper(*args, **kwargs):
        auth = request.headers.get("Authorization", "")
        if not auth.startswith("Bearer ") or auth[7:] != APP_TOKEN:
            return jsonify({"erro": "não autorizado"}), 401
        return f(*args, **kwargs)
    return wrapper

# ── SCRAPING (roda em thread separada) ───────────────────────────────────────

def _rodar_scraper(cookies: list):
    """Executa scraper.py em subprocess e aguarda resultado."""
    estado["scraping_em_curso"] = True
    estado["scraping_erro"]     = None
    log.info("Iniciando scraper em background...")

    raw = ""
    try:
        python = sys.executable
        proc   = subprocess.Popen(
            [python, "-u", SCRAPER_PATH],  # -u = stdout unbuffered
            stdin=subprocess.PIPE,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
        payload = json.dumps({"cookies": cookies}).encode()

        try:
            out, err = proc.communicate(input=payload, timeout=180)
        except subprocess.TimeoutExpired:
            proc.kill()
            out, err = proc.communicate()
            estado["scraping_erro"] = "timeout"
            log.error("Scraper excedeu o tempo limite (180s).")
            return

        # Loga stderr completo de uma vez, sem misturar com outros logs
        if err:
            for linha in err.decode(errors="replace").splitlines():
                log.info(f"[scraper] {linha}")

        raw = out.decode(errors="replace").strip()

        if not raw:
            log.error("Scraper não retornou nenhum JSON.")
            estado["scraping_erro"] = "sem resposta do scraper"
            return

        # Resultado está sempre na última linha do stdout
        ultima_linha = raw.splitlines()[-1]
        resultado    = json.loads(ultima_linha)
        log.info(f"Scraper finalizado: {resultado}")

        if not resultado.get("ok"):
            estado["scraping_erro"] = resultado.get("erro", "erro desconhecido")

        estado["ultimo_scraping"] = int(time.time())

    except json.JSONDecodeError as e:
        estado["scraping_erro"] = f"JSON inválido: {e}"
        log.error(f"Erro ao parsear resultado do scraper: {e} | raw: {raw[:300]}")
    except Exception as e:
        estado["scraping_erro"] = str(e)
        log.error(f"Erro no scraper: {e}")
    finally:
        estado["scraping_em_curso"] = False


def iniciar_scraping_async(cookies: list):
    t = threading.Thread(target=_rodar_scraper, args=(cookies,), daemon=True)
    t.start()

# ── ROTAS DE SISTEMA ──────────────────────────────────────────────────────────

@app.get("/api/health")
def health():
    return jsonify({
        "ok":            True,
        "ts":            int(time.time()),
        "status":        "online",
        "sessao_ativa":  estado["sessao"] is not None,
        "total_projetos": len(estado["projetos"]),
    })

# ── ROTAS DE SESSÃO ───────────────────────────────────────────────────────────

@app.post("/api/workana/sessao")
@requer_auth
def salvar_sessao():
    dados   = request.get_json(silent=True) or {}
    cookies = dados.get("cookies", [])
    auth    = dados.get("auth")

    if not cookies and not auth:
        return jsonify({"erro": "nenhuma sessão recebida"}), 400

    estado["sessao"] = {
        "cookies":  cookies,
        "auth":     auth,
        "salvo_em": int(time.time()),
    }
    _salvar_json("sessao.json", estado["sessao"])
    log.info(f"Sessão Workana salva ({len(cookies)} cookies).")

    # Mantém browser headless aberto para preservar a sessão
    session_keeper.start(cookies)

    # Dispara scraping automático logo após login
    iniciar_scraping_async(cookies)

    return jsonify({"ok": True, "cookies": len(cookies), "scraping": "iniciado"})


@app.get("/api/workana/sessao")
@requer_auth
def obter_sessao():
    sessao = estado["sessao"]
    if not sessao:
        return jsonify({"ativa": False})

    ativa = len(sessao.get("cookies", [])) > 0

    return jsonify({
        "ativa":    ativa,
        "cookies":  len(sessao.get("cookies", [])),
        "salvo_em": sessao.get("salvo_em"),
    })

# ── ROTAS DE PROJETOS ─────────────────────────────────────────────────────────

@app.post("/api/workana/projetos")
@requer_auth
def salvar_projetos():
    dados    = request.get_json(silent=True) or {}
    projetos = dados.get("projetos", [])

    if not isinstance(projetos, list):
        return jsonify({"erro": "formato inválido"}), 400

    estado["projetos"]           = projetos
    estado["projetos_salvos_em"] = int(time.time())

    _salvar_json("projetos.json", {
        "projetos": projetos,
        "total":    len(projetos),
        "salvo_em": estado["projetos_salvos_em"],
    })

    log.info(f"Cache atualizado: {len(projetos)} projetos.")
    return jsonify({"ok": True, "total": len(projetos)})


@app.get("/api/workana/projetos")
@requer_auth
def listar_projetos():
    projetos  = estado["projetos"]
    busca     = request.args.get("busca", "").lower()
    categoria = request.args.get("categoria", "").lower()
    habilidade = request.args.get("habilidade", "").lower()
    tipo      = request.args.get("tipo", "").lower()

    resultado = projetos

    if busca:
        resultado = [p for p in resultado if busca in (p.get("titulo") or "").lower() or busca in (p.get("descricao") or "").lower()]
    if categoria:
        resultado = [p for p in resultado if categoria in (p.get("categoria") or "").lower()]
    if habilidade:
        resultado = [p for p in resultado if any(habilidade in str(h).lower() for h in (p.get("habilidades") or []))]
    if tipo:
        resultado = [p for p in resultado if (p.get("orcamento", {}).get("tipo") or "").lower() == tipo]

    resultado.sort(key=lambda p: p.get("publicado_em") or "", reverse=True)

    return jsonify({
        "ok":           True,
        "total":        len(resultado),
        "projetos":     resultado,
        "cache_age_min": round((time.time() - estado["projetos_salvos_em"]) / 60, 1) if estado["projetos_salvos_em"] else None,
    })

# ── SCRAPING MANUAL ───────────────────────────────────────────────────────────

@app.post("/api/workana/scraping/iniciar")
@requer_auth
def iniciar_scraping():
    if estado["scraping_em_curso"]:
        return jsonify({"ok": False, "erro": "scraping já em andamento"}), 409

    sessao = estado["sessao"]
    if not sessao:
        return jsonify({"ok": False, "erro": "sem sessão — faça login primeiro"}), 400

    cookies = sessao.get("cookies", [])
    iniciar_scraping_async(cookies)

    return jsonify({"ok": True, "status": "iniciado"})


@app.get("/api/workana/scraping/status")
@requer_auth
def status_scraping():
    return jsonify({
        "em_curso":       estado["scraping_em_curso"],
        "ultimo_scraping": estado["ultimo_scraping"],
        "erro":           estado["scraping_erro"],
        "total_projetos": len(estado["projetos"]),
    })

# ── PROJETOS SALVOS ───────────────────────────────────────────────────────────

def _salvar_projetos_salvos():
    _salvar_json("projetos_salvos.json", estado["projetos_salvos"])


@app.post("/api/projetos-salvos")
@requer_auth
def criar_projeto_salvo():
    dados = request.get_json(silent=True) or {}
    oportunidade = dados.get("oportunidade")
    if not oportunidade:
        return jsonify({"erro": "oportunidade obrigatória"}), 400

    # Evita duplicatas pelo id da oportunidade
    op_id = str(oportunidade.get("id", ""))
    for p in estado["projetos_salvos"]:
        if str(p.get("oportunidade", {}).get("id", "")) == op_id and op_id:
            return jsonify({"erro": "projeto já salvo", "projeto": p}), 409

    projeto = {
        "id":            str(uuid.uuid4()),
        "oportunidade":  oportunidade,
        "proposta_texto": "",
        "status":        "em_elaboracao",
        "criado_em":     int(time.time()),
    }
    estado["projetos_salvos"].append(projeto)
    _salvar_projetos_salvos()
    log.info(f"Projeto salvo: {projeto['id']} — {oportunidade.get('titulo', '')}")
    return jsonify(projeto), 201


@app.get("/api/projetos-salvos")
@requer_auth
def listar_projetos_salvos():
    lista = sorted(
        estado["projetos_salvos"],
        key=lambda p: p.get("criado_em", 0),
        reverse=True,
    )
    return jsonify(lista)


@app.delete("/api/projetos-salvos/<projeto_id>")
@requer_auth
def deletar_projeto_salvo(projeto_id):
    original = len(estado["projetos_salvos"])
    estado["projetos_salvos"] = [
        p for p in estado["projetos_salvos"] if p["id"] != projeto_id
    ]
    if len(estado["projetos_salvos"]) == original:
        return jsonify({"erro": "projeto não encontrado"}), 404
    _salvar_projetos_salvos()
    return jsonify({"ok": True})


@app.put("/api/projetos-salvos/<projeto_id>")
@requer_auth
def atualizar_projeto_salvo(projeto_id):
    dados = request.get_json(silent=True) or {}
    for p in estado["projetos_salvos"]:
        if p["id"] == projeto_id:
            if "proposta_texto" in dados:
                p["proposta_texto"] = dados["proposta_texto"]
            if "status" in dados:
                p["status"] = dados["status"]
            _salvar_projetos_salvos()
            return jsonify(p)
    return jsonify({"erro": "projeto não encontrado"}), 404


# ── PROPOSTA IA ────────────────────────────────────────────────────────────────

@app.post("/api/proposta/ia")
@requer_auth
def gerar_proposta_ia():
    dados = request.get_json(silent=True) or {}
    projeto_id = dados.get("projeto_id")
    instrucao  = dados.get("instrucao", "").strip()

    projeto = next(
        (p for p in estado["projetos_salvos"] if p["id"] == projeto_id), None
    )
    if not projeto:
        return jsonify({"erro": "projeto não encontrado"}), 404

    op       = projeto["oportunidade"]
    titulo   = op.get("titulo", "")
    descricao = op.get("descricao", "")
    sobre    = op.get("sobre_projeto", "")
    skills   = ", ".join(op.get("skills", []))
    valor    = op.get("valor", "A combinar")

    proposta_atual = projeto.get("proposta_texto", "")

    instrucao_extra = f"\n\nInstrução adicional do usuário: {instrucao}" if instrucao else ""
    contexto_proposta = (
        f"\n\nTexto atual da proposta (melhore ou reescreva conforme necessário):\n{proposta_atual}"
        if proposta_atual.strip()
        else ""
    )

    prompt = f"""Você é um especialista em propostas comerciais para projetos de tecnologia e TI na plataforma Workana.

Escreva uma proposta profissional em português brasileiro para o seguinte projeto:

Título: {titulo}
Orçamento: {valor}
Habilidades requeridas: {skills}
Descrição do projeto: {descricao}
{f"Sobre o projeto: {sobre}" if sobre else ""}

A proposta deve seguir esta estrutura:
1. Objetivo — entendimento do que o cliente precisa
2. Escopo — o que será entregue concretamente
3. Prazo — estimativa de prazo de entrega
4. Investimento — valor proposto e forma de pagamento
5. Suporte — garantia e suporte pós-entrega
6. Considerações Finais — encerramento cordial e convite para conversa

Tom: profissional, direto, confiante. Não use linguagem genérica. Mostre domínio técnico.{instrucao_extra}{contexto_proposta}

Responda APENAS com o texto da proposta, sem comentários adicionais."""

    provider = os.environ.get("AI_PROVIDER", "anthropic").lower()
    model    = os.environ.get("AI_MODEL", "claude-opus-4-6")

    if provider == "anthropic":
        api_key = os.environ.get("ANTHROPIC_API_KEY", "")
        if not api_key:
            return jsonify({
                "erro": "ANTHROPIC_API_KEY não configurada. Defina no arquivo .env."
            }), 400
        try:
            import anthropic as _anthropic
            client = _anthropic.Anthropic(api_key=api_key)
            msg = client.messages.create(
                model=model,
                max_tokens=2048,
                messages=[{"role": "user", "content": prompt}],
            )
            texto = msg.content[0].text
            return jsonify({"texto": texto})
        except ImportError:
            return jsonify({
                "erro": "Pacote 'anthropic' não instalado. Execute: pip install anthropic"
            }), 500
        except Exception as e:
            return jsonify({"erro": f"Erro Anthropic: {e}"}), 500

    elif provider == "openai":
        api_key = os.environ.get("OPENAI_API_KEY", "")
        if not api_key:
            return jsonify({
                "erro": "OPENAI_API_KEY não configurada. Defina no arquivo .env."
            }), 400
        try:
            import openai as _openai
            client = _openai.OpenAI(api_key=api_key)
            resp = client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=2048,
            )
            texto = resp.choices[0].message.content
            return jsonify({"texto": texto})
        except ImportError:
            return jsonify({
                "erro": "Pacote 'openai' não instalado. Execute: pip install openai"
            }), 500
        except Exception as e:
            return jsonify({"erro": f"Erro OpenAI: {e}"}), 500

    return jsonify({"erro": f"AI_PROVIDER inválido: '{provider}'. Use 'anthropic' ou 'openai'."}), 400


# ── PROPOSTA DOCX ──────────────────────────────────────────────────────────────

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "public", "proposta.docx")


@app.post("/api/proposta/exportar")
@requer_auth
def exportar_proposta():
    dados = request.get_json(silent=True) or {}
    projeto_id = dados.get("projeto_id")

    projeto = next(
        (p for p in estado["projetos_salvos"] if p["id"] == projeto_id), None
    )
    if not projeto:
        return jsonify({"erro": "projeto não encontrado"}), 404

    proposta_texto = projeto.get("proposta_texto", "")
    op = projeto["oportunidade"]
    titulo = op.get("titulo", "Projeto")

    try:
        from docx import Document
        from docx.shared import Pt
        import datetime

        doc = Document(TEMPLATE_PATH)

        # Preenche a tabela de cabeçalho (primeira tabela do documento)
        if doc.tables:
            tabela = doc.tables[0]
            hoje = datetime.date.today().strftime("%d/%m/%Y")
            validade = (datetime.date.today() + datetime.timedelta(days=15)).strftime("%d/%m/%Y")

            for row in tabela.rows:
                for cell in row.cells:
                    txt = cell.text.strip().lower()
                    if "cliente" in txt:
                        cell.paragraphs[0].clear()
                        run = cell.paragraphs[0].add_run(titulo[:60])
                        run.font.size = Pt(10)
                    elif "data" in txt and "validade" not in txt:
                        cell.paragraphs[0].clear()
                        run = cell.paragraphs[0].add_run(hoje)
                        run.font.size = Pt(10)
                    elif "validade" in txt:
                        cell.paragraphs[0].clear()
                        run = cell.paragraphs[0].add_run(validade)
                        run.font.size = Pt(10)
                    elif "proponente" in txt:
                        cell.paragraphs[0].clear()
                        run = cell.paragraphs[0].add_run("Pedro")
                        run.font.size = Pt(10)

        # Remove parágrafos de conteúdo após a tabela do cabeçalho
        # (mantém apenas a tabela; adiciona o texto da proposta)
        # Identifica o índice do primeiro elemento que é a tabela
        body = doc.element.body
        # Remove todos os parágrafos que vêm após a primeira tabela
        tabela_encontrada = False
        elementos_remover = []
        for elemento in body:
            tag = elemento.tag.split("}")[-1] if "}" in elemento.tag else elemento.tag
            if tag == "tbl":
                tabela_encontrada = True
                continue
            if tabela_encontrada and tag == "p":
                elementos_remover.append(elemento)

        for el in elementos_remover:
            body.remove(el)

        # Adiciona o texto da proposta como parágrafos
        for linha in proposta_texto.split("\n"):
            p = doc.add_paragraph(linha)
            p.style.font.size = Pt(10)

        saida = os.path.join(PROPOSTAS_DIR, f"{projeto_id}.docx")
        doc.save(saida)
        log.info(f"Proposta exportada: {saida}")
        return jsonify({"path": f"data/propostas/{projeto_id}.docx"})

    except ImportError:
        return jsonify({
            "erro": "Pacote 'python-docx' não instalado. Execute: pip install python-docx"
        }), 500
    except Exception as e:
        log.error(f"Erro ao exportar proposta: {e}")
        return jsonify({"erro": f"Erro ao gerar DOCX: {e}"}), 500


@app.get("/api/proposta/download/<projeto_id>")
@requer_auth
def baixar_proposta(projeto_id):
    # Valida o projeto_id para evitar path traversal
    projeto_id_limpo = re.sub(r"[^a-zA-Z0-9\-]", "", projeto_id)
    caminho = os.path.join(PROPOSTAS_DIR, f"{projeto_id_limpo}.docx")
    if not os.path.exists(caminho):
        return jsonify({"erro": "arquivo não encontrado — exporte primeiro"}), 404

    projeto = next(
        (p for p in estado["projetos_salvos"] if p["id"] == projeto_id_limpo), None
    )
    titulo_slug = re.sub(r"[^a-zA-Z0-9_\- ]", "", projeto["oportunidade"].get("titulo", "proposta") if projeto else "proposta")[:40]
    nome_arquivo = f"proposta_{titulo_slug}.docx".replace(" ", "_")

    return send_file(
        caminho,
        as_attachment=True,
        download_name=nome_arquivo,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.post("/api/proposta/preview")
@requer_auth
def preview_proposta():
    dados = request.get_json(silent=True) or {}
    projeto_id = dados.get("projeto_id")

    projeto_id_limpo = re.sub(r"[^a-zA-Z0-9\-]", "", projeto_id or "")
    caminho = os.path.join(PROPOSTAS_DIR, f"{projeto_id_limpo}.docx")
    if not os.path.exists(caminho):
        return jsonify({"erro": "arquivo não encontrado — exporte primeiro"}), 404

    try:
        import mammoth
        with open(caminho, "rb") as f:
            resultado = mammoth.convert_to_html(f)
        return jsonify({"html": resultado.value})
    except ImportError:
        return jsonify({
            "erro": "Pacote 'mammoth' não instalado. Execute: pip install mammoth"
        }), 500
    except Exception as e:
        return jsonify({"erro": f"Erro ao converter DOCX: {e}"}), 500


# ── AÇÃO WORKANA ───────────────────────────────────────────────────────────────

_acao_pendente = {}


@app.post("/api/workana/abrir-proposta")
@requer_auth
def abrir_proposta_workana():
    dados = request.get_json(silent=True) or {}
    url   = dados.get("url", "")
    texto = dados.get("texto", "")

    if not url:
        return jsonify({"erro": "url obrigatória"}), 400

    _acao_pendente["url"]   = url
    _acao_pendente["texto"] = texto

    return jsonify({"ok": True, "url": url, "texto": texto})


# ── ROTA PÚBLICA ──────────────────────────────────────────────────────────────

@app.get("/api/oportunidades")
def listar_oportunidades():
    projetos  = estado["projetos"]
    resultado = []

    for i, p in enumerate(projetos):
        orcamento = p.get("orcamento") or {}
        valor_raw = orcamento.get("valor") or orcamento.get("maximo") or orcamento.get("minimo")
        valor     = f"R$ {valor_raw}" if valor_raw else "A combinar"

        resultado.append({
            "id":            p.get("id") or str(i),
            "titulo":        p.get("titulo") or "Sem título",
            "valor":         valor,
            "skills":        p.get("habilidades") or p.get("skills") or [],
            "descricao":     p.get("descricao") or "",
            "sobre_projeto": p.get("sobre_projeto") or "",
            "publicado_em":  p.get("publicado_em") or "",
            "url":           p.get("url") or "",
        })

    resultado.sort(key=lambda p: p.get("publicado_em") or "", reverse=True)
    return jsonify(resultado)

# ── ENTRADA ───────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    log.info(f"Servidor WorkanaBot iniciado na porta {SERVER_PORT}")

    # Se já houver sessão salva, inicia o keeper e o scraping
    if estado["sessao"]:
        cookies_salvos = estado["sessao"].get("cookies", [])
        log.info("Sessão encontrada no disco — iniciando keeper e scraping...")
        session_keeper.start(cookies_salvos)
        iniciar_scraping_async(cookies_salvos)

    app.run(host="127.0.0.1", port=SERVER_PORT, debug=False)